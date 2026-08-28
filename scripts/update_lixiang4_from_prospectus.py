#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按 20260819 募集说明书 Word 更新 lixiang-4：最新数据写入，【】标黄。"""

from __future__ import annotations

import shutil
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))
from apply_lixiang_template import (  # noqa: E402
    AUTHOR,
    BLANK,
    add_table_row,
    enable_track_revisions,
    insert_column,
    revise_cell,
    revise_p,
    revise_tc,
)

SRC = Path("/home/ubuntu/.cursor/projects/workspace/uploads/lixiang-4_ac60.docx")
PROS = Path("/home/ubuntu/.cursor/projects/workspace/uploads/________________2026______________________20260819_649f.docx")
DESKTOP = Path("/home/ubuntu/Desktop")
REPO = Path("/workspace/docs/incoming")
FILE1 = "济宁市兖州区惠民城建投资有限公司非公开发行公司债券项目立项申请报告.docx"


def vis_tc(tc) -> str:
    parts = []
    for t in tc.findall(".//" + qn("w:t")):
        cur = t
        skipped = False
        while cur is not None:
            if cur.tag == qn("w:del"):
                skipped = True
                break
            cur = cur.getparent()
        if not skipped:
            parts.append(t.text or "")
    return "".join(parts).replace("\n", "").strip()


def vis_cell(cell) -> str:
    return vis_tc(cell._tc)


def norm(s: str) -> str:
    return (s or "").replace(" ", "").replace("\u3000", "").strip()


def fill_from_src(dst_table, src_table, header_map):
    """按首列项目名，把 src 各期写入 dst（dst 已含对应列表头）。"""
    src_map = {}
    for tr in src_table._tbl.findall(qn("w:tr"))[1:]:
        tcs = tr.findall(qn("w:tc"))
        if not tcs:
            continue
        key = norm(vis_tc(tcs[0]))
        src_map[key] = [vis_tc(tc) for tc in tcs]
    for tr in dst_table._tbl.findall(qn("w:tr"))[1:]:
        tcs = tr.findall(qn("w:tc"))
        if len(tcs) < 2:
            continue
        key = norm(vis_tc(tcs[0]))
        if key not in src_map:
            continue
        srow = src_map[key]
        for dst_i, src_i in header_map:
            if dst_i < len(tcs) and src_i < len(srow):
                val = srow[src_i]
                if val in ("【】", "【 】"):
                    val = BLANK
                if val == "":
                    continue
                revise_tc(tcs[dst_i], val)


def replace_interest_table(dst, src56):
    """表头改为 2026年6月末 / 2025年末 / 2024年末，按募说表填写。"""
    trs = dst._tbl.findall(qn("w:tr"))
    head = trs[0].findall(qn("w:tc"))
    # 实际4格：项目、一年以内、2026.6、2025 → 改为 2026.6、2025、2024
    revise_tc(head[1], "2026年6月末")
    revise_tc(head[2], "2025年末")
    revise_tc(head[3], "2024年末")
    src_rows = {}
    for tr in src56._tbl.findall(qn("w:tr"))[2:]:
        tcs = tr.findall(qn("w:tc"))
        src_rows[norm(vis_tc(tcs[0]))] = [vis_tc(c) for c in tcs]
    for tr in trs[2:]:
        tcs = tr.findall(qn("w:tc"))
        key = norm(vis_tc(tcs[0]))
        if key not in src_rows:
            continue
        s = src_rows[key]
        # src: 0项目 1-2 2026金额占比 3-4 2025 5-6 2024
        for i in range(1, 7):
            val = s[i] if i < len(s) else ""
            if val in ("【】", "【 】"):
                val = BLANK
            revise_tc(tcs[i], val)


def insert_period_column(table, header: str):
    insert_column(table, 1)
    tr0 = table._tbl.findall(qn("w:tr"))[0]
    tcs = tr0.findall(qn("w:tc"))
    if len(tcs) > 1:
        revise_tc(tcs[1], header)
    # 原第二列表头可能变成空，补回
    if len(tcs) > 2 and not vis_tc(tcs[2]):
        pass


def fill_fs(dst, src, new_header):
    insert_column(dst, 1)
    trs = dst._tbl.findall(qn("w:tr"))
    revise_tc(trs[0].findall(qn("w:tc"))[1], new_header)
    src_map = {}
    for tr in src._tbl.findall(qn("w:tr"))[1:]:
        tcs = tr.findall(qn("w:tc"))
        src_map[norm(vis_tc(tcs[0]))] = vis_tc(tcs[1]) if len(tcs) > 1 else ""
    for tr in trs[1:]:
        tcs = tr.findall(qn("w:tc"))
        key = norm(vis_tc(tcs[0]))
        if key in src_map and len(tcs) > 1:
            val = src_map[key]
            if val in ("【】", "【 】"):
                val = BLANK
            revise_tc(tcs[1], val)


def replace_credit_table(dst, src):
    src_rows = []
    for tr in src._tbl.findall(qn("w:tr")):
        tcs = tr.findall(qn("w:tc"))
        src_rows.append([vis_tc(c) for c in tcs])
    while len(dst.rows) < len(src_rows):
        add_table_row(dst)
    for ri, srow in enumerate(src_rows):
        tcs = dst._tbl.findall(qn("w:tr"))[ri].findall(qn("w:tc"))
        for ci, val in enumerate(srow):
            if ci < len(tcs):
                if val in ("【】", "【 】"):
                    val = BLANK
                revise_tc(tcs[ci], val)


def expand_related(dst, src):
    insert_column(dst, 2)
    trs = dst._tbl.findall(qn("w:tr"))
    revise_tc(trs[0].findall(qn("w:tc"))[2], "2026年6月末")
    # 按 项目名称+关联方 匹配
    src_map = {}
    for tr in src._tbl.findall(qn("w:tr"))[1:]:
        tcs = tr.findall(qn("w:tc"))
        key = (norm(vis_tc(tcs[0])), norm(vis_tc(tcs[1])))
        src_map[key] = vis_tc(tcs[2]) if len(tcs) > 2 else ""
    for tr in trs[1:]:
        tcs = tr.findall(qn("w:tc"))
        key = (norm(vis_tc(tcs[0])), norm(vis_tc(tcs[1])))
        if key in src_map and len(tcs) > 2:
            val = src_map[key]
            if val in ("【】", "【 】", ""):
                val = BLANK
            revise_tc(tcs[2], val)


def expand_income(dst, src):
    """收入表保持 金额/占比 分组，在最前插入 2026年1-6月 两组。"""
    # 插入两列于 index 1
    insert_column(dst, 1)
    insert_column(dst, 1)
    trs = dst._tbl.findall(qn("w:tr"))
    head = trs[0].findall(qn("w:tc"))
    # 表头可能是 3 个实际格（项目 + 两年各 span2）。插入后结构易乱，直接改可见表头文字
    if len(head) >= 2:
        revise_tc(head[1], "2026年1-6月")
    sub = trs[1].findall(qn("w:tc"))
    if len(sub) >= 3:
        revise_tc(sub[1], "金额")
        revise_tc(sub[2], "占比")
    src_map = {}
    for tr in src._tbl.findall(qn("w:tr"))[2:]:
        tcs = tr.findall(qn("w:tc"))
        src_map[norm(vis_tc(tcs[0]))] = (vis_tc(tcs[1]), vis_tc(tcs[2]))
    for tr in trs[2:]:
        tcs = tr.findall(qn("w:tc"))
        key = norm(vis_tc(tcs[0]))
        if key in src_map and len(tcs) >= 3:
            a, b = src_map[key]
            if a in ("【】",):
                a = BLANK
            if b in ("【】",):
                b = BLANK
            if a in ("-", "—"):
                a = "-"
            revise_tc(tcs[1], a)
            revise_tc(tcs[2], b)


def main():
    DESKTOP.mkdir(parents=True, exist_ok=True)
    REPO.mkdir(parents=True, exist_ok=True)

    out_desk = DESKTOP / FILE1
    out4 = DESKTOP / "lixiang-4.docx"
    shutil.copy2(SRC, out_desk)

    doc = Document(str(out_desk))
    enable_track_revisions(doc)
    pros = Document(str(PROS))

    # 有息负债表 → 募说表56
    replace_interest_table(doc.tables[2], pros.tables[56])

    # 授信明细改募说表73，正文时点：2026.6仍空，明细为募说披露的2024年末表
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if "截至2025年末，发行人银行授信情况具体如下" in t or "发行人银行授信情况具体如下" in t:
            # 保留 2026.6 留空句，明细改称 2024 年末（与募说授信表一致）
            vis = "".join(
                (x.text or "")
                for x in p._p.findall(".//" + qn("w:t"))
                if not any(a.tag == qn("w:del") for a in x.iterancestors())
            )
            if "具体如下" in vis or "具体如下" in t:
                new = vis.replace("截至2025年末，发行人银行授信情况具体如下", "截至2024年末，发行人银行授信情况具体如下")
                if new == vis:
                    if "具体如下" in vis and "2024年末" not in vis:
                        new = vis.replace("截至2025年末", "截至2024年末")
                if new != vis and new.strip():
                    revise_p(p, new)
    replace_credit_table(doc.tables[3], pros.tables[73])

    # 三大报表插入 2026 列
    fill_fs(doc.tables[9], pros.tables[24], "2026年6月30日")
    fill_fs(doc.tables[10], pros.tables[25], "2026年1-6月")
    fill_fs(doc.tables[11], pros.tables[26], "2026年1-6月")
    fill_fs(doc.tables[12], pros.tables[27], "2026年6月30日")
    fill_fs(doc.tables[13], pros.tables[28], "2026年1-6月")
    fill_fs(doc.tables[14], pros.tables[29], "2026年1-6月")

    expand_related(doc.tables[6], pros.tables[67])
    expand_income(doc.tables[7], pros.tables[8])

    doc.save(str(out_desk))
    shutil.copy2(out_desk, out4)
    shutil.copy2(out_desk, REPO / FILE1)
    shutil.copy2(out_desk, REPO / "lixiang-4.docx")
    shutil.copy2(out_desk, REPO / "lixiang.docx")
    print("wrote", out_desk)
    print("wrote", out4)


if __name__ == "__main__":
    main()
