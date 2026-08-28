#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""修正立项稿表头串行、授信口径、营收错行、子公司口径，并去掉误标黄。"""

from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from apply_lixiang_template import (  # noqa: E402
    AUTHOR,
    BLANK,
    FILE1,
    _make_run,
    enable_track_revisions,
    revise_p,
    revise_tc,
)

DESKTOP = Path("/home/ubuntu/Desktop")
REPO = Path("/workspace/docs/incoming")
SRC = DESKTOP / "lixiang-4.docx"

INTEREST_ROWS = [
    # 项目, 2026.6金额/占比, 2025金额/占比, 2024金额/占比
    ("银行贷款", BLANK, BLANK, BLANK, BLANK, "50.11", "18.85"),
    ("其中担保贷款", BLANK, BLANK, BLANK, BLANK, "50.11", "18.85"),
    ("其中：政策性银行", BLANK, BLANK, BLANK, BLANK, "7.65", "2.88"),
    ("国有六大行", BLANK, BLANK, BLANK, BLANK, "1.00", "0.38"),
    ("股份制银行", BLANK, BLANK, BLANK, BLANK, "24.87", "9.36"),
    ("地方城商行", BLANK, BLANK, BLANK, BLANK, "5.22", "5.73"),
    ("地方农商行", BLANK, BLANK, BLANK, BLANK, "1.38", "0.52"),
    ("其他银行", BLANK, BLANK, BLANK, BLANK, "-", "-"),
    ("债券融资", BLANK, BLANK, BLANK, BLANK, "164.97", "62.07"),
    ("其中：公司债券", "36.59", "15.60", "38.49", "15.63", "38.49", "14.48"),
    ("企业债券", BLANK, BLANK, BLANK, BLANK, "-", "-"),
    ("债务融资工具", BLANK, BLANK, BLANK, BLANK, "126.48", "47.59"),
    ("非标融资", BLANK, BLANK, BLANK, BLANK, "50.70", "19.08"),
    ("其中：信托融资", BLANK, BLANK, BLANK, BLANK, "2.99", "1.13"),
    ("融资租赁", BLANK, BLANK, BLANK, BLANK, "7.08", "2.66"),
    ("其他非标融资", BLANK, BLANK, BLANK, BLANK, "40.63", "15.29"),
    ("其他融资", BLANK, BLANK, BLANK, BLANK, "-", "-"),
    ("合计", "231.99", "100.00", "246.20", "100.00", "265.78", "100.00"),
]

INCOME_ROWS = [
    ("市政工程业务", "90,742.48", "88.29", "207,154.56", "60.47", "200,200.00", "37.52"),
    ("商品销售", "8,019.50", "7.80", "49,730.46", "14.52", "299,048.06", "56.05"),
    ("房屋销售业务", "3,560.29", "3.46", "7,958.36", "2.32", "18,918.78", "3.55"),
    ("其他业务", "90.73", "0.09", "19.42", "0.01", "23.67", "0.00"),
    ("主营业务收入小计", "102,322.26", "99.56", "264,862.80", "77.31", "518,190.51", "97.12"),
    ("转让土地使用权业务", "-", "-", "57,102.71", "16.67", "4,282.30", "0.80"),
    ("房租业务", "361.92", "0.35", "9,832.35", "2.87", "10,810.62", "2.03"),
    ("其他业务", "-", "-", "10,792.94", "3.15", "265.49", "0.05"),
    ("其他业务收入小计", "361.92", "0.35", "77,728.00", "22.69", "15,358.41", "2.88"),
    ("合计", "102,774.90", "100.00", "342,590.80", "100.00", "533,548.93", "100.00"),
]

SUB_ROWS = [
    ("1", "济宁市惠丰产业发展投资有限公司", "济宁市兖州区", "100,000.00", "100.00", "一级"),
    ("2", "济宁市兖州区国开齐鲁文化旅游开发有限公司", "济宁市兖州区", "38,000.00", "100.00", "一级"),
    ("3", "济宁市裕丰粮食产业有限公司", "济宁市兖州区", "19,600.00", "100.00", "一级"),
    ("4", "山东兴隆文化园有限公司", "济宁市兖州区", "3,000.00", "100.00", "一级"),
    ("5", "济宁市兖州区帅元文化旅游发展有限公司", "济宁市兖州区", "10,000.00", "90.00", "二级"),
    ("6", "济宁市兖州区端信文旅产业发展有限公司", "济宁市兖州区", "300.00", "100.00", "二级"),
    ("7", "济宁市兖州区兴隆寺文化传媒有限公司", "济宁市兖州区", "30.00", "100.00", "二级"),
    ("8", "济宁市兖州区兴隆电子商贸有限公司", "济宁市兖州区", "3.00", "100.00", "二级"),
    ("9", "济宁市兖州区融科供应链管理有限公司", "济宁市兖州区", BLANK, BLANK, BLANK),
    ("10", "山东惠园工投招商服务有限公司", "济宁市兖州区", BLANK, BLANK, BLANK),
]


def vis_tc(tc) -> str:
    parts = []
    for t in tc.findall(".//" + qn("w:t")):
        cur = t
        skipped = False
        while cur is not None:
            if cur.tag == qn("w:del"):
                skipped = True
                break
            cur = cur.getparent()
        if not skipped:
            parts.append(t.text or "")
    return "".join(parts).replace("\n", "").strip()


def first_rpr(tc):
    for r in tc.findall(".//" + qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            pr = deepcopy(rPr)
            hl = pr.find(qn("w:highlight"))
            if hl is not None:
                pr.remove(hl)
            return pr
    return None


def accept_revisions(root) -> None:
    """接受元素内已有修订，消除表头/单元格里叠层的删除+插入串文。"""
    while True:
        nodes = root.findall(".//" + qn("w:del")) + root.findall(".//" + qn("w:moveFrom"))
        if not nodes:
            break
        nodes.sort(key=lambda e: len(list(e.iterancestors())))
        node = nodes[-1]
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)

    while True:
        nodes = root.findall(".//" + qn("w:ins")) + root.findall(".//" + qn("w:moveTo"))
        if not nodes:
            break
        nodes.sort(key=lambda e: len(list(e.iterancestors())))
        node = nodes[-1]
        parent = node.getparent()
        if parent is None:
            break
        idx = list(parent).index(node)
        for child in list(node):
            node.remove(child)
            parent.insert(idx, child)
            idx += 1
        parent.remove(node)


def set_tc_text(tc, text: str) -> None:
    ps = tc.findall(qn("w:p"))
    pPr = deepcopy(ps[0].find(qn("w:pPr"))) if ps and ps[0].find(qn("w:pPr")) is not None else None
    rPr = first_rpr(tc)
    for p in list(ps):
        tc.remove(p)
    p = OxmlElement("w:p")
    if pPr is not None:
        p.append(pPr)
    p.append(_make_run(text, rPr, highlight=(text == BLANK)))
    tc.append(p)


def set_span(tc, n: int) -> None:
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    gs = tcPr.find(qn("w:gridSpan"))
    if n <= 1:
        if gs is not None:
            tcPr.remove(gs)
        return
    if gs is None:
        gs = OxmlElement("w:gridSpan")
        tcPr.append(gs)
    gs.set(qn("w:val"), str(n))


def strip_misplaced_highlight(root) -> int:
    """已有完整数据不得标黄；仅「——」保留黄底。"""
    removed = 0
    for r in root.findall(".//" + qn("w:r")):
        text = "".join((t.text or "") for t in r.findall(qn("w:t")))
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        hl = rPr.find(qn("w:highlight"))
        if hl is None:
            continue
        if text == BLANK:
            continue
        rPr.remove(hl)
        removed += 1
    return removed


def rebuild_period_header(tr, labels: list[str]) -> None:
    """项目 + 三段各占两列。labels 如 [2026年6月末, 2025年末, 2024年末]。"""
    tcs = tr.findall(qn("w:tc"))
    if not tcs:
        return
    set_tc_text(tcs[0], "项目")
    set_span(tcs[0], 1)
    # 需要 1 + 3 = 4 个表头格
    while len(tr.findall(qn("w:tc"))) > 4:
        extra = tr.findall(qn("w:tc"))[-1]
        tr.remove(extra)
    tcs = tr.findall(qn("w:tc"))
    while len(tcs) < 4:
        new_tc = deepcopy(tcs[-1])
        tr.append(new_tc)
        tcs = tr.findall(qn("w:tc"))
    for i, label in enumerate(labels, start=1):
        set_tc_text(tcs[i], label)
        set_span(tcs[i], 2)


def ensure_data_rows(tbl, n_data: int, template_tr):
    trs = tbl.findall(qn("w:tr"))
    # keep header + subheader
    while len(tbl.findall(qn("w:tr"))) > 2 + n_data:
        tbl.remove(tbl.findall(qn("w:tr"))[-1])
    while len(tbl.findall(qn("w:tr"))) < 2 + n_data:
        new_tr = deepcopy(template_tr)
        for t in new_tr.findall(".//" + qn("w:t")):
            t.text = ""
        tbl.append(new_tr)


def fill_7col_row(tr, values: list[str]) -> None:
    tcs = tr.findall(qn("w:tc"))
    while len(tcs) < 7:
        new_tc = deepcopy(tcs[-1])
        tr.append(new_tc)
        tcs = tr.findall(qn("w:tc"))
    while len(tcs) > 7:
        tr.remove(tcs[-1])
        tcs = tr.findall(qn("w:tc"))
    for i, val in enumerate(values):
        set_span(tcs[i], 1)
        set_tc_text(tcs[i], val)


def fix_interest_table(table) -> None:
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    rebuild_period_header(trs[0], ["2026年6月末", "2025年末", "2024年末"])
    sub = trs[1].findall(qn("w:tc"))
    labels = ["", "金额", "占比", "金额", "占比", "金额", "占比"]
    while len(sub) < 7:
        new_tc = deepcopy(sub[-1])
        trs[1].append(new_tc)
        sub = trs[1].findall(qn("w:tc"))
    for i, lab in enumerate(labels):
        set_span(sub[i], 1)
        set_tc_text(sub[i], lab)
    ensure_data_rows(tbl, len(INTEREST_ROWS), trs[2])
    trs = tbl.findall(qn("w:tr"))
    for i, row in enumerate(INTEREST_ROWS):
        fill_7col_row(trs[i + 2], list(row))


def fix_income_table(table) -> None:
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    rebuild_period_header(trs[0], ["2026年1-6月", "2025年度", "2024年度"])
    sub = trs[1].findall(qn("w:tc"))
    labels = ["", "金额", "占比", "金额", "占比", "金额", "占比"]
    while len(sub) < 7:
        new_tc = deepcopy(sub[-1])
        trs[1].append(new_tc)
        sub = trs[1].findall(qn("w:tc"))
    for i, lab in enumerate(labels):
        set_span(sub[i], 1)
        set_tc_text(sub[i], lab)
    ensure_data_rows(tbl, len(INCOME_ROWS), trs[2])
    trs = tbl.findall(qn("w:tr"))
    for i, row in enumerate(INCOME_ROWS):
        fill_7col_row(trs[i + 2], list(row))


def fix_sub_table(table) -> None:
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    headers = ["序号", "名称", "注册地", "注册资本（万元）", "持股比例（%）", "级别"]
    head = trs[0].findall(qn("w:tc"))
    for i, h in enumerate(headers):
        set_tc_text(head[i], h)
    template = trs[1]
    while len(tbl.findall(qn("w:tr"))) > 1 + len(SUB_ROWS):
        tbl.remove(tbl.findall(qn("w:tr"))[-1])
    while len(tbl.findall(qn("w:tr"))) < 1 + len(SUB_ROWS):
        new_tr = deepcopy(template)
        tbl.append(new_tr)
    trs = tbl.findall(qn("w:tr"))
    for i, row in enumerate(SUB_ROWS):
        tcs = trs[i + 1].findall(qn("w:tc"))
        for ci, val in enumerate(row):
            set_tc_text(tcs[ci], val)


def delete_element(elm) -> None:
    parent = elm.getparent()
    if parent is not None:
        parent.remove(elm)


def main() -> None:
    doc = Document(str(SRC))
    body = doc.element.body
    accept_revisions(body)
    n_hl = strip_misplaced_highlight(body)
    enable_track_revisions(doc)

    fix_interest_table(doc.tables[2])
    fix_income_table(doc.tables[7])
    fix_sub_table(doc.tables[8])

    # 授信：只保留 2026 年最新口径，删 2024 明细表和单位行
    credit_only = (
        "发行人资信情况良好，与银行等金融机构一直保持长期合作关系。"
        "截至2026年6月末，发行人在各金融机构的授信额度总额为——万元，"
        "其中已使用额度为——万元，剩余授信余额为——万元。"
    )
    for p in doc.paragraphs:
        vis = "".join(
            (x.text or "")
            for x in p._p.findall(".//" + qn("w:t"))
            if not any(a.tag == qn("w:del") for a in x.iterancestors())
        )
        if "发行人在各金融机构的授信额度总额" in vis and "银行授信情况具体如下" in vis:
            revise_p(p, credit_only)
        if vis.strip() == "单位：万元":
            # 授信表明细上方的单位行；担保表另有自己的单位行
            prev = p._p.getprevious()
            prev_txt = ""
            if prev is not None:
                prev_txt = "".join(prev.itertext())
            if "授信" in prev_txt:
                delete_element(p._p)

    # 删除 2024 年末授信明细表（现 tables[3]）
    credit_tbl = doc.tables[3]._tbl
    delete_element(credit_tbl)

    sub_text = (
        "截至2026年6月末，发行人纳入合并报表范围的子公司为10家。"
        "相比2025年末（9家），发行人划入济宁市兖州区融科供应链管理有限公司、"
        "山东惠园工投招商服务有限公司，济宁市兖州区惠民均和实业发展有限公司不再纳入合并报表范围。"
        "一级子公司仍为4家，与2025年末一致，基本情况如下："
    )
    for p in doc.paragraphs:
        vis = "".join(
            (x.text or "")
            for x in p._p.findall(".//" + qn("w:t"))
            if not any(a.tag == qn("w:del") for a in x.iterancestors())
        )
        if vis.startswith("截至2026年6月末，发行人纳入合并报表范围的子公司为10家"):
            revise_p(p, sub_text)

    huifeng = (
        "截至2025年末，惠丰产业总资产为2,489,181.59万元，净资产为862,384.90万元，"
        "负债合计为1,626,796.69万元。2025年度惠丰产业实现营业收入107,925.54万元，"
        "实现净利润-32,631.76万元。惠丰产业净利润为负主要系财务费用较高所致；"
        "2025年度营业收入下降主要系子公司商品销售收入减少所致。"
    )
    for p in doc.paragraphs:
        vis = "".join(
            (x.text or "")
            for x in p._p.findall(".//" + qn("w:t"))
            if not any(a.tag == qn("w:del") for a in x.iterancestors())
        )
        if vis.startswith("截至2025年末，惠丰产业总资产"):
            revise_p(p, huifeng)

    strip_misplaced_highlight(body)

    out_desk = DESKTOP / FILE1
    out4 = DESKTOP / "lixiang-4.docx"
    doc.save(str(out4))
    shutil.copy2(out4, out_desk)
    REPO.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out4, REPO / FILE1)
    shutil.copy2(out4, REPO / "lixiang-4.docx")
    shutil.copy2(out4, REPO / "lixiang.docx")
    print("stripped_highlight_runs", n_hl)
    print("wrote", out4)


if __name__ == "__main__":
    main()
