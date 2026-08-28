#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""对外担保表按募说最新表替换，个人行格式与表内其他行一致。"""

from __future__ import annotations

import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from apply_lixiang_template import FILE1, enable_track_revisions  # noqa: E402
from fix_lixiang_headers import set_tc_text, vis_tc  # noqa: E402

DESKTOP = Path("/home/ubuntu/Desktop")
REPO = Path("/workspace/docs/incoming")
SRC = DESKTOP / "lixiang-4.docx"
PROS = Path(
    "/home/ubuntu/.cursor/projects/workspace/uploads/"
    "________________2026______________________20260819_649f.docx"
)


def norm_date(s: str) -> str:
    t = (s or "").strip()
    if t in ("/", "—", "——"):
        return "-"
    if t.startswith("最晚"):
        t = t.replace("最晚", "").strip()
    m = re.match(r"^(\d{4})[/-](\d{1,2})[/-](\d{1,2})$", t)
    if m:
        return f"{m.group(1)}/{int(m.group(2)):02d}/{int(m.group(3)):02d}"
    return t


def main() -> None:
    doc = Document(str(SRC))
    pros = Document(str(PROS))
    enable_track_revisions(doc)

    src_tbl = pros.tables[68]._tbl
    dst_tbl = doc.tables[3]._tbl
    src_rows = []
    for tr in src_tbl.findall(qn("w:tr")):
        src_rows.append([vis_tc(tc) for tc in tr.findall(qn("w:tc"))])

    template = dst_tbl.findall(qn("w:tr"))[1]
    while len(dst_tbl.findall(qn("w:tr"))) > len(src_rows):
        dst_tbl.remove(dst_tbl.findall(qn("w:tr"))[-1])
    while len(dst_tbl.findall(qn("w:tr"))) < len(src_rows):
        dst_tbl.append(deepcopy(template))

    trs = dst_tbl.findall(qn("w:tr"))
    for ri, srow in enumerate(src_rows):
        tcs = trs[ri].findall(qn("w:tc"))
        vals = list(srow) + [""] * (len(tcs) - len(srow))
        if ri > 0:
            if len(vals) > 3:
                vals[3] = norm_date(vals[3])
            if len(vals) > 4:
                vals[4] = norm_date(vals[4])
        for ci, val in enumerate(vals[: len(tcs)]):
            set_tc_text(tcs[ci], val)
            # 去掉手动换行，与其他行一样只保留一段
            for br in tcs[ci].findall(".//" + qn("w:br")):
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)

    out4 = DESKTOP / "lixiang-4.docx"
    out1 = DESKTOP / FILE1
    doc.save(str(out4))
    shutil.copy2(out4, out1)
    shutil.copy2(out4, REPO / FILE1)
    shutil.copy2(out4, REPO / "lixiang-4.docx")
    shutil.copy2(out4, REPO / "lixiang.docx")
    print("guarantee rows", len(src_rows))
    print("wrote", out4)


if __name__ == "__main__":
    main()
