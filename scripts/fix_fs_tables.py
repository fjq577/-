#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""三大报表按募说重写，确保 2026 年列可见且不错行。"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from apply_lixiang_template import BLANK, FILE1, enable_track_revisions  # noqa: E402
from fix_lixiang_headers import set_tc_text, vis_tc  # noqa: E402

DESKTOP = Path("/home/ubuntu/Desktop")
REPO = Path("/workspace/docs/incoming")
PROS = Path(
    "/home/ubuntu/.cursor/projects/workspace/uploads/"
    "________________2026______________________20260819_649f.docx"
)

# 立项表序号, 募说表序号, 三列表头
PAIRS = [
    (8, 24, ["项目", "2026年6月30日", "2025年12月31日", "2024年12月31日"]),
    (9, 25, ["项目", "2026年1-6月", "2025年度", "2024年度"]),
    (10, 26, ["项目", "2026年1-6月", "2025年度", "2024年度"]),
    (11, 27, ["项目", "2026年6月30日", "2025年12月31日", "2024年12月31日"]),
    (12, 28, ["项目", "2026年1-6月", "2025年度", "2024年度"]),
    (13, 29, ["项目", "2026年1-6月", "2025年度", "2024年度"]),
]


def equalize_period_cols(tbl) -> None:
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    cols = grid.findall(qn("w:gridCol"))
    if len(cols) < 4:
        return
    # 项目列略宽，三期等宽，避免 2026 列被挤没
    cols[0].set(qn("w:w"), "3600")
    for c in cols[1:]:
        c.set(qn("w:w"), "2400")


def fill_table(dst, src, headers: list[str]) -> None:
    src_rows = []
    for tr in src._tbl.findall(qn("w:tr")):
        src_rows.append([vis_tc(tc) for tc in tr.findall(qn("w:tc"))])
    dst_tbl = dst._tbl
    trs = dst_tbl.findall(qn("w:tr"))
    if len(trs) != len(src_rows):
        raise SystemExit(f"row count mismatch {len(trs)} vs {len(src_rows)}")
    for ri, srow in enumerate(src_rows):
        tcs = trs[ri].findall(qn("w:tc"))
        vals = list(srow)
        if ri == 0:
            vals = headers
        else:
            while len(vals) < len(tcs):
                vals.append("")
            for i, v in enumerate(vals):
                if v in ("【】", "【 】"):
                    vals[i] = BLANK
        for ci, val in enumerate(vals[: len(tcs)]):
            set_tc_text(tcs[ci], val)
    equalize_period_cols(dst_tbl)


def main() -> None:
    doc = Document(str(DESKTOP / "lixiang-4.docx"))
    pros = Document(str(PROS))
    enable_track_revisions(doc)
    for dst_i, src_i, headers in PAIRS:
        fill_table(doc.tables[dst_i], pros.tables[src_i], headers)

    out4 = DESKTOP / "lixiang-4.docx"
    out1 = DESKTOP / FILE1
    doc.save(str(out4))
    shutil.copy2(out4, out1)
    shutil.copy2(out4, REPO / FILE1)
    shutil.copy2(out4, REPO / "lixiang-4.docx")
    shutil.copy2(out4, REPO / "lixiang.docx")
    print("rewrote 6 FS tables")


if __name__ == "__main__":
    main()
