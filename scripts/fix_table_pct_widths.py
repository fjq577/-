#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""表格 pct 列宽合计必须为 5000，否则 Word/WPS 会把多出的 2026 列挤成不可见。"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from apply_lixiang_template import FILE1, enable_track_revisions  # noqa: E402

DESKTOP = Path("/home/ubuntu/Desktop")
REPO = Path("/workspace/docs/incoming")


def set_tc_w(tc, w: int) -> None:
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    el = tcPr.find(qn("w:tcW"))
    if el is None:
        el = OxmlElement("w:tcW")
        tcPr.insert(0, el)
    el.set(qn("w:w"), str(w))
    el.set(qn("w:type"), "pct")


def span_of(tc) -> int:
    gs = tc.find(qn("w:tcPr") + "/" + qn("w:gridSpan"))
    return int(gs.get(qn("w:val"))) if gs is not None else 1


def set_tbl_fixed(tbl) -> None:
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tw = tblPr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
        tblPr.append(tw)
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")


def apply_widths(table, grid_widths: list[int]) -> None:
    tbl = table._tbl
    set_tbl_fixed(tbl)
    for tr in tbl.findall(qn("w:tr")):
        slot = 0
        for tc in tr.findall(qn("w:tc")):
            sp = span_of(tc)
            set_tc_w(tc, sum(grid_widths[slot : slot + sp]))
            slot += sp
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
        if i < len(grid_widths):
            gc.set(qn("w:w"), str(int(grid_widths[i] * 1.8)))


def main() -> None:
    doc = Document(str(DESKTOP / "lixiang-4.docx"))
    enable_track_revisions(doc)
    w4 = [1400, 1200, 1200, 1200]
    for i in (8, 9, 10, 11, 12, 13, 14):
        apply_widths(doc.tables[i], w4)
    apply_widths(doc.tables[5], [800, 1500, 900, 900, 900])
    apply_widths(doc.tables[6], [1400, 600, 600, 600, 600, 600, 600])
    out4 = DESKTOP / "lixiang-4.docx"
    out1 = DESKTOP / FILE1
    doc.save(str(out4))
    shutil.copy2(out4, out1)
    shutil.copy2(out4, REPO / FILE1)
    shutil.copy2(out4, REPO / "lixiang-4.docx")
    shutil.copy2(out4, REPO / "lixiang.docx")
    print("normalized pct widths")


if __name__ == "__main__":
    main()
